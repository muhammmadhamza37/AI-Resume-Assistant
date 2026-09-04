import io
import json
import os
import re

import streamlit as st
from docx import Document
from google import genai
from pypdf import PdfReader


MODEL_NAME = "gemini-3.6-flash"
MAX_TEXT_CHARS = 50000


# ---------------------------------------------------------
# PAGE CONFIGURATION
# ---------------------------------------------------------

st.set_page_config(
    page_title="ATS Resume Analyzer",
    page_icon="📄",
    layout="wide"
)


# ---------------------------------------------------------
# RESUME TEXT EXTRACTION
# ---------------------------------------------------------

def extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF resume."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []

    for page in reader.pages:
        pages.append(page.extract_text() or "")

    return "\n".join(pages).strip()


def extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX resume."""
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
            )
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_resume(uploaded_file) -> str:
    """Detect file type and extract resume text."""
    data = uploaded_file.getvalue()
    extension = uploaded_file.name.lower().rsplit(".", 1)[-1]

    if extension == "pdf":
        text = extract_pdf(data)

    elif extension == "docx":
        text = extract_docx(data)

    elif extension == "txt":
        text = data.decode("utf-8", errors="ignore")

    else:
        raise ValueError(
            "Unsupported file type. Please upload PDF, DOCX, or TXT."
        )

    return text[:MAX_TEXT_CHARS]


# ---------------------------------------------------------
# API KEY
# ---------------------------------------------------------

def get_api_key() -> str:
    """Get Gemini API key from Streamlit secrets or environment variable."""
    try:
        key = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        key = ""

    return key or os.getenv("GEMINI_API_KEY", "")


# ---------------------------------------------------------
# GEMINI JSON CLEANING
# ---------------------------------------------------------

def clean_json(text: str) -> dict:
    """Convert Gemini JSON output into a Python dictionary."""
    text = text.strip()

    text = re.sub(
        r"^```(?:json)?\s*",
        "",
        text,
        flags=re.I
    )

    text = re.sub(
        r"\s*```$",
        "",
        text
    )

    return json.loads(text)


# ---------------------------------------------------------
# GEMINI ANALYSIS
# ---------------------------------------------------------

def analyze_resume(resume_text: str, job_description: str) -> dict:
    api_key = get_api_key()

    if not api_key:
        raise RuntimeError(
            "GEMINI_API_KEY is not configured. "
            "Add it to Streamlit Secrets before running the app."
        )

    client = genai.Client(api_key=api_key)

    job_description = job_description.strip()

    if not job_description:
        job_description = (
            "No job description was provided. "
            "Evaluate general ATS readiness and resume quality."
        )

    prompt = f"""
You are an expert ATS resume evaluator, recruiter, and resume optimization specialist.

Analyze the resume against the provided job description.

IMPORTANT RULES:
1. Do NOT invent experience.
2. Do NOT invent education.
3. Do NOT invent skills.
4. Do NOT invent certifications.
5. Do NOT invent employers.
6. Do NOT invent dates.
7. Do NOT invent achievements.
8. Do NOT falsely claim that a keyword exists.
9. If a metric is needed but not provided, use placeholders such as [X%], [X], or [$X].
10. Give practical, specific, actionable recommendations.
11. Missing keywords must come from important terms actually present in the job description.
12. Return ONLY valid JSON with no markdown before or after it.

Use EXACTLY this JSON structure:

{{
  "ats_score": 0,
  "score_breakdown": {{
    "keyword_match": 0,
    "skills_alignment": 0,
    "experience_relevance": 0,
    "format_ats_readability": 0,
    "impact_and_achievements": 0
  }},
  "summary": "Short overall assessment",
  "matched_keywords": [
    "keyword 1",
    "keyword 2"
  ],
  "missing_keywords": [
    "keyword 1",
    "keyword 2"
  ],
  "critical_issues": [
    "Issue 1",
    "Issue 2"
  ],
  "improvements": [
    {{
      "priority": "High",
      "issue": "Specific issue",
      "recommendation": "Specific recommendation",
      "example": "Example improvement"
    }}
  ],
  "section_feedback": {{
    "summary": "Feedback about professional summary",
    "experience": "Feedback about work experience",
    "skills": "Feedback about skills",
    "education": "Feedback about education",
    "projects": "Feedback about projects"
  }},
  "ats_checklist": {{
    "standard_headings": true,
    "simple_format": true,
    "contact_information_present": true,
    "action_verbs": true,
    "quantified_achievements": true,
    "keyword_alignment": true
  }}
}}

SCORING SYSTEM:
Overall ATS score = 0 to 100

Keyword Match = 0-25
Skills Alignment = 0-20
Experience Relevance = 0-20
ATS Format Readability = 0-20
Impact & Achievements = 0-15

If no actual job description is provided:
- Evaluate general ATS readiness.
- Do NOT pretend it is a job-specific match.
- Mention this limitation in the summary.

JOB DESCRIPTION:
{job_description}

RESUME:
{resume_text}
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt
    )

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")

    return clean_json(response.text)


# ---------------------------------------------------------
# UI
# ---------------------------------------------------------

st.title("📄 ATS Resume Analyzer")

st.caption(
    "Upload your resume and get an ATS score, keyword analysis, "
    "section feedback, and practical resume improvements using Gemini."
)


with st.sidebar:
    st.header("How It Works")

    st.markdown(
        """
1. Upload your resume
2. Paste the job description
3. Click **Analyze Resume**
4. Review your ATS score
5. Apply the recommended improvements
"""
    )

    st.divider()

    st.info(
        "For the most useful ATS score, paste the job description "
        "for the role you are targeting."
    )


uploaded = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx", "txt"]
)

job_description = st.text_area(
    "Job Description (Recommended)",
    height=220,
    placeholder=(
        "Paste the job description here to get a job-specific ATS score..."
    )
)


if uploaded:
    try:
        resume_text = extract_resume(uploaded)

    except Exception as exc:
        st.error(f"Could not read the resume: {exc}")
        st.stop()

    if not resume_text:
        st.error("No readable text was found in the resume.")
        st.info(
            "If your PDF is scanned or image-only, upload a text-based "
            "PDF, DOCX, or TXT file."
        )
        st.stop()

    with st.expander("👀 Preview Extracted Resume Text"):
        st.text(resume_text[:10000])

    if st.button(
        "🔍 Analyze Resume",
        type="primary",
        use_container_width=True
    ):
        with st.spinner("Analyzing your resume with Gemini..."):
            try:
                result = analyze_resume(
                    resume_text=resume_text,
                    job_description=job_description
                )

            except json.JSONDecodeError:
                st.error(
                    "Gemini returned an unexpected response format. "
                    "Please try again."
                )
                st.stop()

            except Exception as exc:
                st.error(f"Analysis failed: {exc}")
                st.stop()

        try:
            score = int(result.get("ats_score", 0))
        except (TypeError, ValueError):
            score = 0

        score = max(0, min(100, score))

        st.subheader("🎯 ATS Score")
        st.progress(score / 100)
        st.metric("Overall ATS Score", f"{score}/100")

        summary = result.get("summary", "")
        if summary:
            st.write(summary)

        st.subheader("📊 Score Breakdown")

        breakdown = result.get("score_breakdown", {})
        cols = st.columns(5)

        breakdown_items = [
            ("Keyword Match", "keyword_match", 25),
            ("Skills", "skills_alignment", 20),
            ("Experience", "experience_relevance", 20),
            ("ATS Format", "format_ats_readability", 20),
            ("Impact", "impact_and_achievements", 15),
        ]

        for col, (label, key, maximum) in zip(cols, breakdown_items):
            value = breakdown.get(key, 0)
            col.metric(label, f"{value}/{maximum}")

        left, right = st.columns(2)

        with left:
            st.subheader("✅ Matched Keywords")
            matched = result.get("matched_keywords", [])

            if matched:
                for keyword in matched:
                    st.write(f"• {keyword}")
            else:
                st.write("No matched keywords identified.")

        with right:
            st.subheader("⚠️ Missing Keywords")
            missing = result.get("missing_keywords", [])

            if missing:
                for keyword in missing:
                    st.write(f"• {keyword}")
            else:
                st.write("No important missing keywords identified.")

        st.subheader("🚨 Critical Issues")

        issues = result.get("critical_issues", [])

        if issues:
            for issue in issues:
                st.warning(issue)
        else:
            st.success("No critical issues identified.")

        st.subheader("💡 Recommended Improvements")

        improvements = result.get("improvements", [])

        if improvements:
            for index, item in enumerate(improvements, start=1):
                priority = item.get("priority", "Medium")
                issue = item.get("issue", "Improvement")
                recommendation = item.get("recommendation", "")
                example = item.get("example", "")

                with st.container(border=True):
                    st.markdown(
                        f"### {index}. {issue}"
                    )
                    st.write(f"**Priority:** {priority}")

                    if recommendation:
                        st.write(recommendation)

                    if example:
                        st.caption(f"Example: {example}")
        else:
            st.success("No major improvements were identified.")

        st.subheader("📌 Section-by-Section Feedback")

        section_feedback = result.get("section_feedback", {})

        if section_feedback:
            for section, feedback in section_feedback.items():
                with st.expander(
                    section.replace("_", " ").title()
                ):
                    st.write(feedback)
        else:
            st.write("No section feedback returned.")

        st.subheader("🧾 ATS Checklist")

        checklist = result.get("ats_checklist", {})

        if checklist:
            for item, passed in checklist.items():
                icon = "✅" if passed else "❌"
                st.write(
                    f"{icon} {item.replace('_', ' ').title()}"
                )
        else:
            st.write("No checklist data returned.")

        st.download_button(
            "⬇️ Download Analysis (JSON)",
            data=json.dumps(
                result,
                indent=2,
                ensure_ascii=False
            ),
            file_name="ats_resume_analysis.json",
            mime="application/json",
            use_container_width=True
        )

else:
    st.info("Upload a resume to begin.")
